"""Resume DOCX template export helpers."""

from __future__ import annotations

import io
import re
from typing import Any


def _add_rich_text(paragraph: Any, text: str) -> None:
    parts = re.split(r"(<b>.*?</b>)", text or "", flags=re.IGNORECASE | re.DOTALL)
    for part in parts:
        if not part:
            continue
        match = re.match(r"<b>(.*?)</b>", part, flags=re.IGNORECASE | re.DOTALL)
        if match:
            run = paragraph.add_run(match.group(1))
            run.bold = True
        else:
            paragraph.add_run(re.sub(r"<[^>]+>", "", part))


def _profile_text(tdata: dict) -> str:
    profile = tdata.get("profile")
    if not isinstance(profile, dict):
        return profile or ""
    pool = profile.get("pool") or []
    default_id = profile.get("default")
    return next(
        (item.get("text", "") for item in pool if item.get("id") == default_id),
        (pool[0].get("text", "") if pool else ""),
    )


def build_template_docx(tdata: dict) -> bytes:
    from docx import Document

    doc = Document()

    basics = tdata.get("basics", {})
    doc.add_heading(basics.get("name") or "简历", level=0)
    contact = " · ".join(
        value
        for value in [
            basics.get("phone"),
            basics.get("email"),
            basics.get("city"),
            basics.get("target_role"),
        ]
        if value
    )
    if contact:
        doc.add_paragraph(contact)

    profile = _profile_text(tdata)
    if profile:
        doc.add_heading("个人总结", level=1)
        _add_rich_text(doc.add_paragraph(), profile)

    for section_key, title in (("projects", "项目经历"), ("internships", "实习经历")):
        items = tdata.get(section_key) or []
        if not items:
            continue
        doc.add_heading(title, level=1)
        for item in items:
            head = f"{item.get('company', '')} — {item.get('role', '')} · {item.get('date', '')}"
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(head)
            run.bold = True
            for bullet in item.get("bullets") or []:
                bullet_paragraph = doc.add_paragraph(style="List Bullet")
                _add_rich_text(bullet_paragraph, bullet)

    skills = tdata.get("skills") or []
    if skills:
        doc.add_heading("技能证书", level=1)
        for skill in skills:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"{skill.get('label', '')}：")
            run.bold = True
            paragraph.add_run(skill.get("text", ""))

    education = tdata.get("education") or []
    if education:
        doc.add_heading("教育背景", level=1)
        for item in education:
            school = item.get("school") or item.get("company") or ""
            major = item.get("major") or item.get("role") or ""
            date = item.get("date", "")
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"{school} · {major} · {date}")
            run.bold = True
            for bullet in item.get("bullets") or []:
                bullet_paragraph = doc.add_paragraph(style="List Bullet")
                _add_rich_text(bullet_paragraph, bullet)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
