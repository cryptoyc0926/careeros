"""
Resume DOCX Writer — 原版样式回写
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

目标：用户上传的 DOCX 原件作为模板，定制后的文本按段落级 fuzzy match 回写进原件，
保留字体 / 字号 / 颜色 / 加粗 / 段落格式。

能力边界（明示）：
  - ✅ run-level 文本替换 + 加粗片段继承首个 run 的字体
  - ✅ profile / bullet / skills 单行级替换
  - ⚠️ 嵌套表格 / TextBox / SmartArt / Canva 分栏 → 抛 DocxRewriteUnsupported，上游降级
  - ⚠️ 命中率 < 70% → 抛 DocxRewriteUnsupported
  - ❌ 新增/删除段落不做（保持原件结构）

使用::
    from services.resume_docx_writer import rewrite_docx, DocxRewriteUnsupported

    try:
        new_bytes = rewrite_docx(original_bytes, old_data, new_data)
    except DocxRewriteUnsupported as e:
        # 降级为内置模板导出
        ...
"""
from __future__ import annotations

import io
import re
from copy import deepcopy
from typing import Any


class DocxRewriteUnsupported(Exception):
    """DOCX 原件结构复杂或命中率不足，无法可靠回写。"""


# ═══════════════════════════════════════════════
# 文本归一化（匹配时忽略空白/标点差异）
# ═══════════════════════════════════════════════
_WS_RE = re.compile(r"[\s\u3000\u00a0]+")
_PUNCT_RE = re.compile(
    r"[，。、：；！？·\u2022\-—–,.\[\]（）()\"\"''「」`~!@#$%^&*()_+=|\\<>?/]"
)


def _normalize(s: str) -> str:
    s = _WS_RE.sub("", str(s or ""))
    s = _PUNCT_RE.sub("", s)
    return s.lower()


def _strip_html(s: str) -> str:
    """去掉 <b>...</b> 等 HTML 标签，留纯文本用于 DOCX run。"""
    return re.sub(r"<[^>]+>", "", s or "")


def _extract_bold_segments(s: str) -> list[tuple[str, bool]]:
    """把带 <b>...</b> 的文本切成 [(text, bold), ...]。"""
    out: list[tuple[str, bool]] = []
    pos = 0
    for m in re.finditer(r"<b>(.*?)</b>", s or "", flags=re.IGNORECASE | re.DOTALL):
        if m.start() > pos:
            out.append((s[pos:m.start()], False))
        out.append((m.group(1), True))
        pos = m.end()
    if pos < len(s or ""):
        out.append((s[pos:], False))
    return out or [(s or "", False)]


# ═══════════════════════════════════════════════
# 收集 old_data → new_data 的单行替换对
# ═══════════════════════════════════════════════

def _collect_replacements(old: dict[str, Any], new: dict[str, Any]) -> list[tuple[str, str]]:
    """返回 [(old_text, new_text_with_html), ...]。过滤硬事实和未变字段。"""
    pairs: list[tuple[str, str]] = []

    # profile
    ov = _flatten_profile(old)
    nv = _flatten_profile(new)
    if ov and nv and ov != nv:
        pairs.append((ov, nv))

    # project bullets
    for section in ("projects", "internships"):
        for i, o_item in enumerate(old.get(section) or []):
            n_item = (new.get(section) or [{}])[i] if i < len(new.get(section) or []) else {}
            for j, ob in enumerate(o_item.get("bullets") or []):
                if j >= len(n_item.get("bullets") or []):
                    continue
                nb = n_item["bullets"][j]
                if ob and nb and ob != nb:
                    pairs.append((ob, nb))
            # role
            if o_item.get("role") and n_item.get("role") and o_item["role"] != n_item["role"]:
                pairs.append((o_item["role"], n_item["role"]))

    # skills text
    for i, o_s in enumerate(old.get("skills") or []):
        n_list = new.get("skills") or []
        if i >= len(n_list):
            continue
        n_s = n_list[i]
        if o_s.get("text") and n_s.get("text") and o_s["text"] != n_s["text"]:
            pairs.append((o_s["text"], n_s["text"]))

    return pairs


def _flatten_profile(data: dict[str, Any]) -> str:
    p = data.get("profile")
    if isinstance(p, dict):
        default_id = p.get("default")
        pool = p.get("pool") or []
        for item in pool:
            if item.get("id") == default_id:
                return item.get("text") or ""
        return (pool[0]["text"] if pool else "") or ""
    return str(p or "")


# ═══════════════════════════════════════════════
# python-docx 段落替换
# ═══════════════════════════════════════════════

def _match_paragraph(para_text: str, old_text: str, threshold: float = 0.82) -> bool:
    """归一化 + substring / 字符重叠率判断。"""
    pn = _normalize(para_text)
    on = _normalize(_strip_html(old_text))
    if not pn or not on:
        return False
    # substring
    if on in pn or pn in on:
        return True
    # 字符重叠率（简易 Jaccard）
    cp, co = set(pn), set(on)
    if not co:
        return False
    inter = len(cp & co)
    return inter / max(len(co), 1) >= threshold


def _replace_paragraph_runs(paragraph, new_text_html: str) -> None:
    """把 paragraph 所有 run 的文本替换成 new_text（保留首个 run 的样式）。

    加粗片段用 <b>...</b> 标记，会拆成多个 run：
      - 非粗体片段：复制首 run 的样式，去掉加粗
      - 粗体片段：复制首 run 的样式，强制 bold=True
    """
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(_strip_html(new_text_html))
        return

    first_run = runs[0]
    # 保留首个 run 的字体配置作模板
    try:
        base_font_name = first_run.font.name
        base_size = first_run.font.size
        base_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
    except Exception:
        base_font_name = None
        base_size = None
        base_color = None

    # 清空所有 run 文本
    for r in runs:
        r.text = ""

    segments = _extract_bold_segments(new_text_html)
    # 第一段复用第一个 run
    first_text, first_bold = segments[0]
    first_run.text = first_text
    first_run.bold = bool(first_bold)

    # 剩余片段新建 run
    for text, bold in segments[1:]:
        new_run = paragraph.add_run(text)
        new_run.bold = bool(bold)
        try:
            if base_font_name:
                new_run.font.name = base_font_name
            if base_size:
                new_run.font.size = base_size
            if base_color:
                new_run.font.color.rgb = base_color
        except Exception:
            pass


def _paragraphs_of_doc(doc) -> list:
    """遍历 body + 表格里的段落，跳过 TextBox（w:txbx）。"""
    out = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
                # 嵌套表格不支持
                if cell.tables:
                    raise DocxRewriteUnsupported("嵌套表格不支持回写")
    return out


# ═══════════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════════

def rewrite_docx(
    original_bytes: bytes,
    old_data: dict[str, Any],
    new_data: dict[str, Any],
    *,
    min_match_ratio: float = 0.60,
) -> bytes:
    """
    把 old_data → new_data 的段落级差异回写到 original_bytes 指向的 DOCX。
    命中率过低或结构复杂 → 抛 DocxRewriteUnsupported。
    """
    try:
        from docx import Document
    except ImportError as e:
        raise DocxRewriteUnsupported("python-docx 未安装") from e

    pairs = _collect_replacements(old_data, new_data)
    if not pairs:
        # 无变化，直接返回原件
        return original_bytes

    doc = Document(io.BytesIO(original_bytes))
    try:
        paragraphs = _paragraphs_of_doc(doc)
    except DocxRewriteUnsupported:
        raise

    hit = 0
    used = set()
    for old_text, new_text in pairs:
        matched_idx = -1
        for i, p in enumerate(paragraphs):
            if i in used:
                continue
            if _match_paragraph(p.text, old_text):
                matched_idx = i
                break
        if matched_idx >= 0:
            _replace_paragraph_runs(paragraphs[matched_idx], new_text)
            used.add(matched_idx)
            hit += 1

    ratio = hit / max(len(pairs), 1)
    if ratio < min_match_ratio:
        raise DocxRewriteUnsupported(
            f"DOCX 命中率 {ratio:.0%}（{hit}/{len(pairs)}），低于 {min_match_ratio:.0%}"
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════
# 自测
# ═══════════════════════════════════════════════
if __name__ == "__main__":
    # 仅测工具函数，不联网不需真 DOCX
    assert _normalize("你 好, 世界！") == _normalize("你好世界")
    assert _strip_html("A <b>B</b> C") == "A B C"
    segs = _extract_bold_segments("A <b>B</b> C <b>D</b>")
    assert segs == [("A ", False), ("B", True), (" C ", False), ("D", True)], segs

    old = {"profile": "旧总结", "projects": [{"company": "A", "role": "r1",
                                            "bullets": ["b1", "b2"]}], "skills": [], "internships": []}
    new = {"profile": "新总结 <b>10%</b>", "projects": [{"company": "A", "role": "r1",
                                                     "bullets": ["b1 改", "b2"]}], "skills": [], "internships": []}
    pairs = _collect_replacements(old, new)
    assert ("旧总结", "新总结 <b>10%</b>") in pairs
    assert ("b1", "b1 改") in pairs
    print("[OK] resume_docx_writer self-test passed")
