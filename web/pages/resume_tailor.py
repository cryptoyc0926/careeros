"""
简历在线编辑 & JD 定制
~~~~~~~~~~~~~~~~~~~~~~~~~~~
三栏交互：
  左：JD 输入 / 历史版本
  中：分段编辑器（profile / projects / internships / skills）
  右：实时 PDF 预览 + 下载

数据流：
  resume_master → tailor(jd) → content_json(session) → renderer → PDF bytes
"""

from __future__ import annotations

import hashlib
import difflib
import html
import json
import sqlite3
from pathlib import Path

import streamlit as st

from config import settings
from services import resume_renderer, resume_tailor
from services.ai_engine import AIError
from components import resume_canvas
from components.ui import page_header, section_title, divider, score_hero_card, alert_success, alert_info, alert_warning, alert_danger

DB_PATH = settings.db_full_path

page_header("简历定制 & 在线编辑")
# 压缩顶部留白
st.markdown(
    """
    <style>
    div.block-container,
    .main .block-container,
    [data-testid="stMainBlockContainer"]{
      padding-top:1.2rem !important;
      max-width:1500px !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

if not settings.has_anthropic_key:
    alert_danger("Claude API Key 未配置，请前往 **设置** 页面。")
    st.stop()


def _db_columns(table: str) -> set[str]:
    conn = sqlite3.connect(DB_PATH)
    try:
        rows = conn.execute(f"PRAGMA table_info({table})").fetchall()
        return {r[1] for r in rows}
    finally:
        conn.close()


def _has_db_column(table: str, column: str) -> bool:
    return column in _db_columns(table)


# ── 读 master ────────────────────────────────────────────
def load_master() -> dict | None:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT * FROM resume_master ORDER BY id LIMIT 1").fetchone()
    conn.close()
    if not row:
        return None
    keys = row.keys() if hasattr(row, "keys") else []
    return {
        "id": row["id"],
        "updated_at": row["updated_at"],
        "basics": json.loads(row["basics_json"]),
        "profile": json.loads(row["profile_json"]),
        "projects": json.loads(row["projects_json"]),
        "internships": json.loads(row["internships_json"]),
        "skills": json.loads(row["skills_json"]),
        "education": json.loads(row["education_json"]),
        "original_docx_blob": row["original_docx_blob"] if "original_docx_blob" in keys else None,
        "original_docx_filename": row["original_docx_filename"] if "original_docx_filename" in keys else None,
    }


def demo_master_fallback() -> dict:
    """公开 Demo 的兜底主简历：只在云端样例数据为空/占位符时使用。"""
    return {
        "id": 0,
        "updated_at": "demo-fallback",
        "basics": {
            "name": "杨 超",
            "phone": "186-8795-0926",
            "email": "bc1chao0926@gmail.com",
            "target_role": "AI 增长运营",
            "city": "杭州",
            "availability": "下周到岗",
            "photo": "",
        },
        "profile": {
            "pool": [
                {
                    "id": "growth_ai",
                    "tags": ["增长", "AI", "运营"],
                    "text": (
                        "统计科班出身，擅长把数据变成增长动作。曾从 0 搭建 AI Trading 社区，"
                        "实现 9,000+ X 粉丝与 1,300+ Telegram 订阅，完成 20+ 次商业合作；"
                        "3 段运营实习覆盖增长、内容与数据分析。"
                    ),
                }
            ],
            "default": "growth_ai",
        },
        "projects": [
            {
                "company": "AI Trading 社区搭建",
                "role": "用户增长运营",
                "date": "2024.03 - 至今",
                "bullets": [
                    "从 0 搭建 X 与 Telegram 内容矩阵，在零投放预算下增长至 9,000+ 粉丝与 1,300+ 订阅。",
                    "设计热点监测、信号筛选、结构化分析、分发输出流程，稳定支撑每周 15+ 条内容产出。",
                    "围绕用户关注点做内容测试与社群运营，沉淀可复用的 AI 内容生产 SOP。",
                ],
            },
            {
                "company": "CareerOS 求职系统",
                "role": "产品与自动化实践",
                "date": "2025.10 - 至今",
                "bullets": [
                    "搭建岗位池、JD 解析、简历定制、投递看板与面试准备工作流，覆盖求职全链路。",
                    "用 SQLite 管理岗位与投递数据，结合大模型完成 JD 匹配、简历改写与评估。",
                ],
            },
        ],
        "internships": [
            {
                "company": "Fancy Tech",
                "role": "海外产品运营实习生",
                "date": "2024.06 - 2024.09",
                "bullets": [
                    "搭建 TikTok/Instagram 内容矩阵，建立 AI 内容生产、自动化编辑、多平台分发 SOP。",
                    "拆解 PhotoRoom、Pebblely 等竞品链路，提炼高转化场景并输出产品优化建议。",
                ],
            },
            {
                "company": "杭银消费金融股份有限公司",
                "role": "数据运营实习生",
                "date": "2023.06 - 2023.10",
                "bullets": [
                    "参与用户分层、活动复盘与指标看板维护，支持运营策略迭代。",
                    "整理业务数据与活动结果，输出周报和专项分析材料。",
                ],
            },
        ],
        "skills": [
            {"label": "数据分析", "text": "Python / SQL / Excel / 指标拆解 / A/B 测试"},
            {"label": "增长运营", "text": "内容矩阵搭建 / 社群运营 / 用户转化 / SOP 沉淀"},
            {"label": "AI 工具", "text": "Prompt Engineering / 自动化工作流 / LLM 应用原型"},
        ],
        "education": [
            {
                "school": "云南财经大学",
                "major": "统计学",
                "date": "2022.09 - 2026.07",
                "bullets": ["核心课程：统计学、计量经济学、Python 数据分析、数据库原理。"],
            }
        ],
    }


def master_needs_demo_fallback(master: dict) -> bool:
    def _bad_text(value: object) -> bool:
        text = str(value or "").strip()
        compact = text.replace("*", "").replace("-", "").replace("—", "").strip()
        return not compact or text.startswith("****")

    basics = master.get("basics", {})
    if _bad_text(basics.get("name")) or str(basics.get("name", "")).startswith("你的"):
        return True
    for section in ("projects", "internships"):
        items = master.get(section) or []
        if not items:
            return True
        for item in items:
            if _bad_text(item.get("company")) or _bad_text(item.get("role")):
                return True
    return False


def flatten_master_for_render(master: dict) -> dict:
    """把 profile pool 折叠成默认文本，得到 renderer 直接可用的扁平结构。"""
    profile = master["profile"]
    if isinstance(profile, dict) and "pool" in profile:
        default_id = profile.get("default")
        text = next(
            (p["text"] for p in profile["pool"] if p["id"] == default_id),
            profile["pool"][0]["text"],
        )
    else:
        text = profile
    return {
        "basics": master["basics"],
        "profile": text,
        "projects": master["projects"],
        "internships": master["internships"],
        "skills": master["skills"],
        "education": master["education"],
    }


master = load_master()
if not master:
    alert_info("主简历还没创建。请先到左侧「**主简历**」页面填写基本信息和经历，再回来使用定制功能。")
    st.stop()
demo_fallback_active = False
if settings.demo_mode and master_needs_demo_fallback(master):
    master = demo_master_fallback()
    demo_fallback_active = True


# ── Session state ────────────────────────────────────────
master_signature = f"{master.get('id')}:{master.get('updated_at', '')}"
if (
    st.session_state.get("_tailor_master_signature") != master_signature
    or
    "tailor_data" not in st.session_state
    or (demo_fallback_active and master_needs_demo_fallback(st.session_state.tailor_data))
):
    st.session_state.tailor_data = flatten_master_for_render(master)
    st.session_state.tailor_meta = {}
    st.session_state.tailor_jd = ""
    st.session_state["_tailor_preview_key"] = None
    st.session_state["_tailor_preview_pdf"] = None
    st.session_state["_tailor_preview_png"] = None
    st.session_state["_tailor_preview_backend"] = None
    st.session_state["_tailor_preview_error"] = None
    st.session_state["_tailor_master_signature"] = master_signature
if "tailor_meta" not in st.session_state:
    st.session_state.tailor_meta = {}
if "tailor_jd" not in st.session_state:
    st.session_state.tailor_jd = ""
if "tailor_chat" not in st.session_state:
    st.session_state.tailor_chat = {"messages": [], "pending": None}
if "tailor_undo_stack" not in st.session_state:
    st.session_state.tailor_undo_stack = []

ENABLE_CHAT_TAILOR = settings.enable_chat_tailor


# ── 撤销栈工具 ────────────────────────────────────────────
import copy as _copy_tailor  # noqa: E402  — 避免顶部再加一个 import copy
_UNDO_LIMIT = 10


def _push_undo_snapshot(label: str = "") -> None:
    """把当前 tailor_data 深拷贝入栈（用于撤销）。"""
    st.session_state.tailor_undo_stack.append(
        {"label": label, "data": _copy_tailor.deepcopy(st.session_state.tailor_data)}
    )
    if len(st.session_state.tailor_undo_stack) > _UNDO_LIMIT:
        st.session_state.tailor_undo_stack = st.session_state.tailor_undo_stack[-_UNDO_LIMIT:]


def _pop_undo() -> str | None:
    if not st.session_state.tailor_undo_stack:
        return None
    snap = st.session_state.tailor_undo_stack.pop()
    st.session_state.tailor_data = snap["data"]
    return snap.get("label") or "上一步"


def _chat_transcript_payload() -> str:
    chat = st.session_state.get("tailor_chat") or {"messages": [], "pending": None}
    payload = {
        "messages": chat.get("messages", []),
        "pending": None,
    }
    return json.dumps(payload, ensure_ascii=False)


def _restore_chat_transcript(raw: str | None) -> None:
    if not raw:
        st.session_state.tailor_chat = {"messages": [], "pending": None}
        return
    try:
        parsed = json.loads(raw)
    except Exception:
        st.session_state.tailor_chat = {"messages": [], "pending": None}
        return
    if isinstance(parsed, list):
        st.session_state.tailor_chat = {"messages": parsed, "pending": None}
    elif isinstance(parsed, dict):
        st.session_state.tailor_chat = {
            "messages": parsed.get("messages", []) or [],
            "pending": None,
        }
    else:
        st.session_state.tailor_chat = {"messages": [], "pending": None}


def _clear_tailor_preview_cache() -> None:
    st.session_state.pop("_tailor_preview_key", None)
    st.session_state.pop("_tailor_preview_pdf", None)
    st.session_state.pop("_tailor_preview_png", None)
    st.session_state.pop("_tailor_preview_backend", None)
    st.session_state.pop("_tailor_preview_error", None)


def _sync_bound_widget(key: str, current_value: object) -> str:
    current = "" if current_value is None else str(current_value)
    source_key = f"_tailor_widget_source::{key}"
    previous_source = st.session_state.get(source_key)
    widget_value = st.session_state.get(key)

    if key not in st.session_state:
        st.session_state[key] = current
    elif previous_source != current and widget_value == previous_source:
        st.session_state[key] = current
    return current


def _bound_text_input(label: str, current_value: object, key: str, **kwargs) -> str:
    current = _sync_bound_widget(key, current_value)
    value = st.text_input(label, key=key, **kwargs)
    if value != current:
        _clear_tailor_preview_cache()
    st.session_state[f"_tailor_widget_source::{key}"] = value
    return value


def _bound_text_area(label: str, current_value: object, key: str, **kwargs) -> str:
    current = _sync_bound_widget(key, current_value)
    value = st.text_area(label, key=key, **kwargs)
    if value != current:
        _clear_tailor_preview_cache()
    st.session_state[f"_tailor_widget_source::{key}"] = value
    return value


# ── Chat 面板（Phase 1 · 全宽折叠区，置于三栏之上）────────
from services import resume_chat as _resume_chat_service  # noqa: E402

from services import resume_patch as _resume_patch_mod  # noqa: E402


def _apply_resume_patch(patch: list[dict], label: str) -> bool:
    """Apply manual/chat edits through the same patch + validation + undo path."""
    new_data, errors = _resume_patch_mod.apply_patch(st.session_state.tailor_data, patch)
    if errors:
        alert_danger("Patch 应用失败：\n" + "\n".join(errors))
        return False
    from services.resume_validator import validate_tailored
    current_report = None
    try:
        current_report = validate_tailored(
            st.session_state.tailor_data,
            st.session_state.tailor_data,
            None,
        )
        report = validate_tailored(st.session_state.tailor_data, new_data, None)
    except Exception:
        report = None

    def _issue_value(issue, field: str) -> str:
        if isinstance(issue, dict):
            return str(issue.get(field, ""))
        return str(getattr(issue, field, ""))

    def _issue_fingerprint(issue) -> tuple[str, str, str, str, str]:
        return (
            _issue_value(issue, "rule"),
            _issue_value(issue, "location"),
            _issue_value(issue, "message"),
            _issue_value(issue, "expected"),
            _issue_value(issue, "actual"),
        )

    baseline_errors = {
        _issue_fingerprint(e)
        for e in (current_report.hard_errors if current_report else [])
    }
    new_hard_errors = [
        e
        for e in (report.hard_errors if report else [])
        if _issue_fingerprint(e) not in baseline_errors
    ]

    if new_hard_errors:
        alert_danger(
            f"硬规则校验失败（新增 {len(new_hard_errors)} 条），拒绝应用：\n"
            + "\n".join(
                f"- [{_issue_value(e, 'rule')}] {_issue_value(e, 'location')}: {_issue_value(e, 'message')}"
                for e in new_hard_errors[:5]
            )
        )
        return False
    _push_undo_snapshot(label=label)
    st.session_state.tailor_data = new_data
    _clear_tailor_preview_cache()
    return True


def _render_patch_diff(patch: list[dict], tailor_data: dict) -> None:
    """把 pending_patch 渲染成字符级 old/new 对比。"""
    st.markdown(
        """
        <style>
        .cos-diff-box{border:1px solid rgba(29,29,31,0.08);border-radius:10px;
          background:#ffffff;margin:8px 0 14px 0;overflow:hidden;}
        .cos-diff-path{padding:8px 12px;background:#f5f5f7;color:#6e6e73;
          font-size:12px;border-bottom:1px solid rgba(29,29,31,0.06);}
        .cos-diff-line{display:flex;gap:10px;padding:10px 12px;white-space:pre-wrap;
          line-height:1.55;font-family:"SF Mono",ui-monospace,Menlo,monospace;
          font-size:13px;color:#1d1d1f;}
        .cos-diff-line + .cos-diff-line{border-top:1px solid rgba(29,29,31,0.06);}
        .cos-diff-prefix{width:16px;flex:0 0 16px;font-weight:700;}
        .cos-diff-minus{color:#c42323;}
        .cos-diff-plus{color:#18794e;}
        .cos-diff-del{background:#ffe8e8;color:#9f1c1c;text-decoration:line-through;}
        .cos-diff-add{background:#dcfce7;color:#11643a;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    def _char_diff_html(old_text: str, new_text: str) -> tuple[str, str]:
        old_parts: list[str] = []
        new_parts: list[str] = []
        for token in difflib.ndiff(list(old_text), list(new_text)):
            tag = token[:1]
            char = html.escape(token[2:])
            if tag == " ":
                old_parts.append(char)
                new_parts.append(char)
            elif tag == "-":
                old_parts.append(f'<span class="cos-diff-del">{char}</span>')
            elif tag == "+":
                new_parts.append(f'<span class="cos-diff-add">{char}</span>')
        return "".join(old_parts), "".join(new_parts)

    for i, p in enumerate(patch):
        path = p.get("path", "")
        new_val = p.get("value", "")
        try:
            old_val = _resume_patch_mod.get_by_path(tailor_data, path)
        except Exception:
            old_val = "(不存在)"
        old_html, new_html = _char_diff_html(str(old_val), str(new_val))
        st.markdown(
            f'''
            <div class="cos-diff-box">
              <div class="cos-diff-path">#{i + 1} · {html.escape(path)}</div>
              <div class="cos-diff-line"><span class="cos-diff-prefix cos-diff-minus">-</span><span>{old_html}</span></div>
              <div class="cos-diff-line"><span class="cos-diff-prefix cos-diff-plus">+</span><span>{new_html}</span></div>
            </div>
            ''',
            unsafe_allow_html=True,
        )


def _render_chat_panel() -> None:
    """Chat 面板：消息流 + pending patch diff + 输入框 + 撤销/清空。"""
    top_c1, top_c2 = st.columns([1, 1])
    with top_c1:
        st.caption(f"撤销栈：{len(st.session_state.tailor_undo_stack)} 步")
    with top_c2:
        if st.button("撤销上一步", key="tailor_undo_btn", use_container_width=True,
                     disabled=not st.session_state.tailor_undo_stack):
            label = _pop_undo()
            if label:
                alert_info(f"已撤销：{label}")
                st.rerun()

    # 用浅色自定义气泡代替 st.chat_message（避免深色头像框）
    st.markdown(
        """
        <style>
        .cos-chat-msg{padding:10px 14px;border-radius:12px;margin:6px 0;
          border:1px solid rgba(29,29,31,0.08);line-height:1.55;}
        .cos-chat-user{background:#f5f7fb;}
        .cos-chat-bot{background:#ffffff;}
        .cos-chat-role{font-size:11px;color:#6e6e73;margin-bottom:4px;letter-spacing:0.3px;}
        [data-testid="stChatInput"] textarea{
          background:#ffffff !important;
          color:#1d1d1f !important;
          border:1px solid rgba(29,29,31,0.10) !important;
        }
        [data-testid="stChatInput"] textarea::placeholder{color:#6e6e73 !important;}
        </style>
        """,
        unsafe_allow_html=True,
    )
    for msg in st.session_state.tailor_chat["messages"][-8:]:
        role = msg["role"]
        klass = "cos-chat-user" if role == "user" else "cos-chat-bot"
        label = "你" if role == "user" else "AI"
        content_html = msg.get("content", "")
        st.markdown(
            f'<div class="cos-chat-msg {klass}"><div class="cos-chat-role">{label}</div>{content_html}</div>',
            unsafe_allow_html=True,
        )
        meta = msg.get("_meta") or {}
        if meta.get("applied"):
            st.caption("已应用到简历 · 可点「撤销上一步」回退")
        if meta.get("advice_md"):
            st.markdown(meta["advice_md"], unsafe_allow_html=True)
        if meta.get("clarify_question"):
            alert_warning(meta["clarify_question"])

    # Pending patch 渲染
    pending = st.session_state.tailor_chat.get("pending")
    if pending and pending.get("patch"):
        alert_info(pending.get("explanation") or "准备修改（待确认）")
        _render_patch_diff(pending["patch"], st.session_state.tailor_data)
        pcol_a, pcol_b = st.columns(2)
        with pcol_a:
            if st.button("应用修改", type="primary",
                         key="tailor_patch_apply",
                         use_container_width=True):
                patch = pending["patch"]
                if _apply_resume_patch(patch, label="chat · 应用 patch"):
                    for m in reversed(st.session_state.tailor_chat["messages"]):
                        if m["role"] == "assistant":
                            m.setdefault("_meta", {})["applied"] = True
                            break
                    st.session_state.tailor_chat["pending"] = None
                    alert_success("已应用，预览可刷新查看")
                    st.rerun()
        with pcol_b:
            if st.button("取消", key="tailor_patch_cancel",
                         use_container_width=True):
                st.session_state.tailor_chat["pending"] = None
                st.rerun()

    user_msg = st.chat_input("告诉 AI 你想怎么改，或问它建议", key="tailor_chat_input")
    if user_msg:
        if st.session_state.tailor_chat.get("pending"):
            st.session_state.tailor_chat["pending"] = None
        st.session_state.tailor_chat["messages"].append(
            {"role": "user", "content": user_msg}
        )
        with st.spinner("AI 思考中..."):
            resp = _resume_chat_service.handle_user_message(
                user_msg=user_msg,
                tailor_data=st.session_state.tailor_data,
                master=master,
                jd_text=st.session_state.tailor_jd or "",
            )
        assistant_meta: dict = {"intent": resp.get("intent")}
        if resp["intent"] == "full_rewrite" and resp.get("new_data"):
            _push_undo_snapshot(label="chat · 整体重写")
            new_data = resp["new_data"]
            new_meta = new_data.pop("_meta", {}) if isinstance(new_data, dict) else {}
            st.session_state.tailor_data = new_data
            st.session_state.tailor_meta = new_meta
            st.session_state.tailor_chat["pending"] = None
            assistant_meta["applied"] = True
            content = resp.get("explanation") or "已整体重写"
        elif resp["intent"] in ("section_rewrite", "patch_ops") and resp.get("pending_patch"):
            st.session_state.tailor_chat["pending"] = {
                "explanation": resp.get("explanation") or "",
                "patch": resp.get("pending_patch"),
            }
            content = resp.get("explanation") or "准备修改（待你确认）"
        elif resp["intent"] == "advice_only":
            assistant_meta["advice_md"] = resp.get("advice_md") or ""
            content = resp.get("explanation") or "建议如下："
        elif resp["intent"] == "clarify":
            assistant_meta["clarify_question"] = resp.get("clarify_question") or ""
            content = resp.get("explanation") or "需要先确认一下："
        else:
            content = f"出错了：{resp.get('error') or resp.get('explanation')}"
        st.session_state.tailor_chat["messages"].append(
            {"role": "assistant", "content": content, "_meta": assistant_meta}
        )
        st.rerun()

    if st.button("清空对话", key="tailor_chat_clear_btn", use_container_width=True,
                 disabled=not st.session_state.tailor_chat["messages"]):
        st.session_state.tailor_chat = {"messages": [], "pending": None}
        st.rerun()


def _current_preview_key() -> str:
    return hashlib.sha256(
        json.dumps(st.session_state.tailor_data, ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()


def _cached_preview() -> tuple[bytes | None, bytes | None, str | None]:
    preview_key = _current_preview_key()
    if st.session_state.get("_tailor_preview_key") != preview_key:
        return None, None, None
    return (
        st.session_state.get("_tailor_preview_pdf"),
        st.session_state.get("_tailor_preview_png"),
        st.session_state.get("_tailor_preview_error"),
    )


def _ensure_preview(force: bool = False) -> tuple[bytes | None, bytes | None, str | None]:
    if not force:
        pdf_bytes, png_bytes, preview_error = _cached_preview()
        if pdf_bytes is not None or preview_error:
            return pdf_bytes, png_bytes, preview_error
    preview_key = _current_preview_key()
    try:
        with st.spinner("正在生成 PDF 预览..."):
            pdf_bytes = resume_renderer.render_pdf_bytes(st.session_state.tailor_data)
            png_bytes, backend = resume_renderer.render_preview_png(
                st.session_state.tailor_data,
                pdf_bytes=pdf_bytes,
            )
        st.session_state["_tailor_preview_key"] = preview_key
        st.session_state["_tailor_preview_pdf"] = pdf_bytes
        st.session_state["_tailor_preview_png"] = png_bytes
        st.session_state["_tailor_preview_backend"] = backend
        st.session_state["_tailor_preview_error"] = None
        return pdf_bytes, png_bytes, None
    except Exception as e:
        preview_error = str(e)
        st.session_state["_tailor_preview_key"] = preview_key
        st.session_state["_tailor_preview_pdf"] = None
        st.session_state["_tailor_preview_png"] = None
        st.session_state["_tailor_preview_backend"] = None
        st.session_state["_tailor_preview_error"] = preview_error
        return None, None, preview_error


def _build_template_docx(tdata: dict) -> bytes:
    from docx import Document
    import io as _io
    import re as _re

    doc = Document()
    b = tdata.get("basics", {})
    doc.add_heading(b.get("name") or "简历", level=0)
    contact = " · ".join(
        x for x in [b.get("phone"), b.get("email"), b.get("city"), b.get("target_role")] if x
    )
    if contact:
        doc.add_paragraph(contact)

    def _add_rich(p, text: str):
        parts = _re.split(r"(<b>.*?</b>)", text or "", flags=_re.IGNORECASE | _re.DOTALL)
        for part in parts:
            if not part:
                continue
            m = _re.match(r"<b>(.*?)</b>", part, flags=_re.IGNORECASE | _re.DOTALL)
            if m:
                r = p.add_run(m.group(1))
                r.bold = True
            else:
                p.add_run(_re.sub(r"<[^>]+>", "", part))

    prof = tdata.get("profile")
    if isinstance(prof, dict):
        pool = prof.get("pool") or []
        default_id = prof.get("default")
        prof = next(
            (x.get("text", "") for x in pool if x.get("id") == default_id),
            (pool[0].get("text", "") if pool else ""),
        )
    if prof:
        doc.add_heading("个人总结", level=1)
        _add_rich(doc.add_paragraph(), prof)

    for sec, title in (("projects", "项目经历"), ("internships", "实习经历")):
        items = tdata.get(sec) or []
        if not items:
            continue
        doc.add_heading(title, level=1)
        for it in items:
            head = f"{it.get('company','')} — {it.get('role','')} · {it.get('date','')}"
            p = doc.add_paragraph()
            r = p.add_run(head)
            r.bold = True
            for bu in (it.get("bullets") or []):
                pb = doc.add_paragraph(style="List Bullet")
                _add_rich(pb, bu)

    skills = tdata.get("skills") or []
    if skills:
        doc.add_heading("技能证书", level=1)
        for s in skills:
            p = doc.add_paragraph()
            r = p.add_run(f"{s.get('label','')}：")
            r.bold = True
            p.add_run(s.get("text", ""))

    edu = tdata.get("education") or []
    if edu:
        doc.add_heading("教育背景", level=1)
        for e in edu:
            school = e.get("school") or e.get("company") or ""
            major = e.get("major") or e.get("role") or ""
            date = e.get("date", "")
            p = doc.add_paragraph()
            r = p.add_run(f"{school} · {major} · {date}")
            r.bold = True
            for bu in (e.get("bullets") or []):
                pb = doc.add_paragraph(style="List Bullet")
                _add_rich(pb, bu)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_save_version() -> None:
    v_name = st.text_input("版本命名", value="", placeholder="如：MiniMax-增长", key="tailor_version_name")
    if st.button("保存此版本", use_container_width=True):
        if not v_name.strip():
            alert_warning("请输入版本名")
            return
        conn = sqlite3.connect(DB_PATH)
        if _has_db_column("resume_versions", "chat_transcript_json"):
            conn.execute(
                """INSERT INTO resume_versions
                   (master_id, target_role, version_name, content_json, match_score, chat_transcript_json)
                   VALUES (?, ?, ?, ?, ?, ?)""",
                (
                    master["id"],
                    st.session_state.tailor_data["basics"].get("target_role"),
                    v_name.strip(),
                    json.dumps(st.session_state.tailor_data, ensure_ascii=False),
                    st.session_state.tailor_meta.get("match_score", 0),
                    _chat_transcript_payload(),
                ),
            )
        else:
            conn.execute(
                """INSERT INTO resume_versions
                   (master_id, target_role, version_name, content_json, match_score)
                   VALUES (?, ?, ?, ?, ?)""",
                (
                    master["id"],
                    st.session_state.tailor_data["basics"].get("target_role"),
                    v_name.strip(),
                    json.dumps(st.session_state.tailor_data, ensure_ascii=False),
                    st.session_state.tailor_meta.get("match_score", 0),
                ),
            )
        conn.commit()
        conn.close()
        alert_success(f"已保存：{v_name}")
        st.rerun()


def _render_history_versions() -> None:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    versions = conn.execute(
        """SELECT id, version_name, target_role, match_score, created_at
           FROM resume_versions ORDER BY created_at DESC LIMIT 10"""
    ).fetchall()
    conn.close()
    if not versions:
        st.caption("暂无")
        return
    for v in versions:
        if st.button(
            f"#{v['id']} {v['version_name']} · {v['match_score']}分",
            key=f"v_{v['id']}",
            use_container_width=True,
        ):
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            has_chat_col = _has_db_column("resume_versions", "chat_transcript_json")
            select_cols = "content_json, chat_transcript_json" if has_chat_col else "content_json"
            row = conn.execute(
                f"SELECT {select_cols} FROM resume_versions WHERE id = ?",
                (v["id"],),
            ).fetchone()
            conn.close()
            st.session_state.tailor_data = json.loads(row["content_json"])
            _restore_chat_transcript(row["chat_transcript_json"] if has_chat_col else None)
            st.session_state.tailor_undo_stack = []
            st.session_state.tailor_meta = {}
            _clear_tailor_preview_cache()
            st.rerun()


def _render_pdf_preview_block(*, thumbnail: bool) -> None:
    if st.button("生成预览", key="generate_preview_left" if thumbnail else "generate_preview_full",
                 use_container_width=True):
        _ensure_preview(force=True)
    pdf_bytes, png_bytes, preview_error = _cached_preview()
    if preview_error:
        alert_danger(f"预览渲染失败：{preview_error}")
    elif png_bytes:
        if thumbnail:
            st.markdown('<div class="cos-preview-thumb">', unsafe_allow_html=True)
            st.image(png_bytes, width=180)
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.image(png_bytes, width=650)
    elif pdf_bytes and not png_bytes:
        import base64 as _b64p
        _b64_pdf = _b64p.b64encode(pdf_bytes).decode()
        st.markdown(
            f'<iframe src="data:application/pdf;base64,{_b64_pdf}" '
            f'width="100%" height="640px" '
            f'style="border:1px solid rgba(29,29,31,0.08);border-radius:14px"></iframe>',
            unsafe_allow_html=True,
        )
    else:
        st.caption("生成预览后显示缩略图" if thumbnail else "生成预览后显示 PDF 大图")


def _render_export_controls() -> None:
    pdf_bytes, _png_bytes, _preview_error = _cached_preview()
    base_name = (
        f"{st.session_state.tailor_data['basics'].get('name', 'resume')}_简历_"
        f"{st.session_state.tailor_data['basics'].get('target_role', '定制版')}"
    )
    st.download_button(
        "PDF",
        data=pdf_bytes or b"",
        file_name=f"{base_name}.pdf",
        mime="application/pdf",
        use_container_width=True,
        disabled=pdf_bytes is None,
        type="primary",
    )

    orig_blob = master.get("original_docx_blob")
    if orig_blob:
        if st.button("DOCX回写", key="dl_docx_inplace", use_container_width=True):
            try:
                from services.resume_docx_writer import rewrite_docx
                old_for_diff = flatten_master_for_render(master)
                new_bytes = rewrite_docx(
                    bytes(orig_blob), old_for_diff, st.session_state.tailor_data
                )
                st.session_state["_last_inplace_docx"] = new_bytes
                alert_success("原版回写成功，点下方确认下载")
            except Exception as err:
                try:
                    fallback = _build_template_docx(st.session_state.tailor_data)
                    st.session_state["_last_inplace_docx"] = fallback
                    alert_warning(f"原版 DOCX 无法回写（{type(err).__name__}），已降级为模板 DOCX")
                except Exception as fallback_err:
                    alert_danger(f"DOCX 生成失败：{fallback_err}")
        if st.session_state.get("_last_inplace_docx"):
            st.download_button(
                "确认下载DOCX",
                data=st.session_state["_last_inplace_docx"],
                file_name=f"{base_name}_原版回写.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_docx_inplace_confirm",
            )
    else:
        st.caption("未上传 DOCX，原版回写不可用")

    try:
        tmpl_docx = _build_template_docx(st.session_state.tailor_data)
        st.download_button(
            "DOCX模板",
            data=tmpl_docx,
            file_name=f"{base_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_docx_template",
        )
    except Exception as e:
        alert_danger(f"DOCX 模板生成失败：{e}")


# ── 布局：两栏（左 JD+Chat / 右 预览+评估）──────────────
col_left, col_right = st.columns([0.9, 2.4])

# ═══ 左栏：JD 输入 + 历史版本 ═══
with col_left:
    st.markdown("##### 目标 JD")

    # 岗位池下拉 —— 关联 job_descriptions，有 JD 的自动带入
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    jobs = conn.execute(
        '''SELECT id, "公司" AS company, "岗位名称" AS position,
                  "等级" AS level, link_type, 链接 AS url,
                  jd_fetch_mode, jd_status, jd_last_error
           FROM jobs_pool ORDER BY id DESC LIMIT 80'''
    ).fetchall()

    # 预先构建 jd_map: 优先用 notes 中的 岗位池#id 精确关联，fallback 到 (company,title) 模糊匹配
    jd_map: dict[int, str] = {}
    all_jds = conn.execute(
        "SELECT company, title, raw_text, notes FROM job_descriptions WHERE raw_text IS NOT NULL"
    ).fetchall()

    # 建 notes 索引: pool_id -> raw_text
    notes_index: dict[int, str] = {}
    for jd in all_jds:
        notes = jd["notes"] or ""
        import re as _re
        m = _re.search(r"岗位池#(\d+)", notes)
        if m:
            notes_index[int(m.group(1))] = jd["raw_text"] or ""

    for j in jobs:
        # 1. notes ID 精确匹配
        txt = notes_index.get(j["id"], "")
        # 2. fallback: company + title 模糊匹配
        if not txt:
            for jd in all_jds:
                if (jd["company"] or "").strip() == (j["company"] or "").strip() and \
                   (jd["title"] or "").strip() == (j["position"] or "").strip():
                    txt = jd["raw_text"] or ""
                    break
        if txt and len(txt.strip()) >= 200 and not txt.startswith("Company:"):
            jd_map[j["id"]] = txt.strip()
    conn.close()

    # 五档状态：已抓 / 待Chrome / 可自动 / 手动 / 黑名单
    def _icon(j) -> str:
        if j["id"] in jd_map:
            return "[已抓]"
        mode = j["jd_fetch_mode"] or "auto"
        status = j["jd_status"] or "pending"
        if mode == "blocked":
            return "[黑名单]"
        if mode == "manual":
            return "[手动]"
        if mode == "browser":
            return "[待Chrome]"
        if mode == "auto":
            return "[失败]" if status == "failed" else "[待抓]"
        return "[?]"

    def _fmt_job(j) -> str:
        return f"{_icon(j)} #{j['id']} [{j['level'] or '-'}] {j['company']} / {j['position']}"

    job_options = ["（手动粘贴 JD）"] + [_fmt_job(j) for j in jobs]
    job_id_by_label = {_fmt_job(j): j["id"] for j in jobs}
    job_meta_by_id = {j["id"]: j for j in jobs}

    choice = st.selectbox(
        "从岗位池选（状态已标注：已抓 / 可自动 / 待Chrome / 手动 / 黑名单）",
        job_options,
        key="job_choice",
    )

    prev = st.session_state.get("_last_job_choice")
    if choice != prev:
        st.session_state._last_job_choice = choice
        if choice == "（手动粘贴 JD）":
            pass
        else:
            jid = job_id_by_label.get(choice)
            if jid in jd_map:
                st.session_state.tailor_jd = jd_map[jid]
                st.rerun()
            else:
                st.session_state.tailor_jd = ""
                st.rerun()

    # 当前选中岗位的状态提示
    if choice != "（手动粘贴 JD）":
        jid = job_id_by_label.get(choice)
        jmeta = job_meta_by_id.get(jid)
        if jmeta is None:
            pass
        elif jid in jd_map:
            alert_success(f"已关联 JD（{len(jd_map[jid])} 字，adapter 抓取）")
        else:
            mode = jmeta["jd_fetch_mode"] or "auto"
            if mode == "blocked":
                alert_danger("黑名单公司（字节/蚂蚁/腾讯/网易），不建议投递")
            elif mode == "manual":
                alert_warning(
                    "此岗位需手动获取 JD\n\n"
                    "邮箱投递 / 纯门户列表页 / 链接失效。到原链接复制 JD 后粘贴到下方。"
                )
            elif mode == "browser":
                alert_info(
                    "此岗位需 Chrome MCP 抓取\n\n"
                    f"原链接：{jmeta['url']}\n\n"
                    "SPA 页面 Python 抓不到。方式二选一：\n"
                    "① 在 Claude 对话里说「用 Chrome MCP 抓岗位池 #"
                    f"{jid} 的 JD」让我来抓；② 手动复制粘贴到下方。"
                )
            elif mode == "auto":
                err = jmeta["jd_last_error"] or ""
                if err:
                    alert_danger(f"自动抓取失败：{err[:120]}")
                else:
                    alert_info("此岗位在自动抓取队列，尚未处理。可到终端跑 `python scripts/jd_worker.py run`")

    jd_text = st.text_area(
        "JD 原文",
        value=st.session_state.tailor_jd,
        height=220,
        placeholder="粘贴 JD 全文...",
        key="jd_textarea",
    )

    col_a, col_b = st.columns(2)
    with col_a:
        go = st.button("生成定制版", type="primary", use_container_width=True)
    with col_b:
        reset = st.button("重置为主简历", use_container_width=True)

    if reset:
        _push_undo_snapshot(label="重置为主简历")
        st.session_state.tailor_data = flatten_master_for_render(master)
        st.session_state.tailor_meta = {}
        st.session_state.tailor_jd = ""
        st.rerun()

    if go and jd_text.strip():
        with st.spinner("正在分析 JD 并重写简历..."):
            try:
                _push_undo_snapshot(label="JD 定制重写")
                tailored = resume_tailor.tailor_resume(master, jd_text)
                meta = tailored.pop("_meta", {})
                st.session_state.tailor_data = tailored
                st.session_state.tailor_meta = meta
                st.session_state.tailor_jd = jd_text
                # 显示校验警告（硬错会进 except 分支）
                v = meta.get("validation", {})
                warn_count = len(v.get("warnings", []))
                if warn_count:
                    alert_warning(f"完成 · 匹配度 {meta.get('match_score', '?')} · {warn_count} 条校验警告（见右栏）")
                else:
                    alert_success(f"完成 · 匹配度 {meta.get('match_score', '?')} · 校验全通过")
                st.rerun()
            except AIError as e:
                alert_danger(f"AI 调用失败：{e}")
            except Exception as e:
                # T6 ValidationError 走这里，渲染红色 banner + diff
                from services.resume_validator import ValidationError
                if isinstance(e, ValidationError):
                    st.session_state.tailor_validation_error = e.report.as_dict()
                    alert_danger(f"硬规则校验失败：{len(e.report.hard_errors)} 条硬错 · 未写入定制版")
                    with st.expander("查看违规明细", expanded=True):
                        for err in e.report.hard_errors:
                            st.markdown(
                                f"- **[{err['rule']}]** `{err['location']}` — {err['message']}\n\n"
                                f"  期望：`{err['expected']}`\n\n  实际：`{err['actual']}`"
                            )
                        if e.report.warnings:
                            st.markdown("---")
                            st.markdown("**警告**：")
                            for w in e.report.warnings:
                                st.caption(f"[{w['rule']}] {w['location']} — {w['message']}")
                else:
                    alert_danger(f"失败：{e}")

    # ── Chat 面板（紧跟在 JD 下方）──
    if ENABLE_CHAT_TAILOR:
        st.markdown("---")
        st.markdown("##### AI 对话修改")
        st.caption("整体重写 / 段落重写 / 精细 patch / 建议 / 反问")
        _render_chat_panel()

    st.markdown("---")
    with st.expander("历史版本", expanded=False):
        _render_history_versions()

    st.markdown("---")
    st.markdown("##### PDF 预览")
    _render_pdf_preview_block(thumbnail=True)

    st.markdown("---")
    st.markdown("##### 导出")
    _render_export_controls()
    _render_save_version()


def _path_state_key(prefix: str, path: str) -> str:
    digest = hashlib.sha1(path.encode("utf-8")).hexdigest()[:12]
    return f"{prefix}_{digest}"


def _open_inline_edit(path: str, current_value: object) -> None:
    st.session_state[_path_state_key("edit_mode", path)] = True
    st.session_state[_path_state_key("edit_draft", path)] = "" if current_value is None else str(current_value)


def _close_inline_edit(path: str) -> None:
    st.session_state.pop(_path_state_key("edit_mode", path), None)
    st.session_state.pop(_path_state_key("edit_draft", path), None)


def _render_edit_pane(path: str, current_value: object, *, height: int = 90, label: str = "内容") -> None:
    draft_key = _path_state_key("edit_draft", path)
    if draft_key not in st.session_state:
        st.session_state[draft_key] = "" if current_value is None else str(current_value)
    st.markdown('<div class="cv-edit-pane">', unsafe_allow_html=True)
    new_value = st.text_area(label, key=draft_key, height=height, label_visibility="collapsed")
    save_col, cancel_col, _ = st.columns([2, 2, 1])
    with save_col:
        if st.button("保存", key=_path_state_key("save", path), use_container_width=True):
            if _apply_resume_patch(
                [{"op": "replace", "path": path, "value": new_value}],
                label=f"手动编辑 {path}",
            ):
                _close_inline_edit(path)
                st.rerun()
    with cancel_col:
        if st.button("取消", key=_path_state_key("cancel", path), use_container_width=True):
            _close_inline_edit(path)
            st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)


def _render_editable_block(path: str, value: object, *, block_class: str, height: int = 90) -> None:
    if st.session_state.get(_path_state_key("edit_mode", path)):
        _render_edit_pane(path, value, height=height)
        return
    text_col, edit_col = st.columns([8, 1])
    with text_col:
        st.markdown(
            f'<div class="{block_class}">{resume_canvas.rich_text_html(value)}</div>',
            unsafe_allow_html=True,
        )
    with edit_col:
        if st.button("编辑", key=_path_state_key("edit", path), use_container_width=True):
            _open_inline_edit(path, value)
            st.rerun()


def _render_basics_editor(data: dict) -> None:
    basics = data.setdefault("basics", {})
    fields = [
        ("name", "姓名"),
        ("email", "邮箱"),
        ("phone", "电话"),
        ("target_role", "求职意向"),
        ("city", "城市"),
        ("availability", "到岗时间"),
    ]
    with st.expander("编辑基本信息", expanded=False):
        cols = st.columns(2)
        values: dict[str, str] = {}
        for idx, (field, label) in enumerate(fields):
            key = f"basics_edit_{field}"
            current = _sync_bound_widget(key, basics.get(field, ""))
            with cols[idx % 2]:
                values[field] = st.text_input(label, key=key, value=current)
                st.session_state[f"_tailor_widget_source::{key}"] = values[field]
        save_col, cancel_col = st.columns(2)
        with save_col:
            if st.button("保存基本信息", key="save_basics_canvas", type="primary", use_container_width=True):
                patch = [
                    {"op": "replace", "path": f"basics.{field}", "value": value}
                    for field, value in values.items()
                ]
                if _apply_resume_patch(patch, label="手动编辑基本信息"):
                    st.rerun()
        with cancel_col:
            if st.button("取消基本信息编辑", key="cancel_basics_canvas", use_container_width=True):
                for field, _label in fields:
                    st.session_state.pop(f"basics_edit_{field}", None)
                    st.session_state.pop(f"_tailor_widget_source::basics_edit_{field}", None)
                st.rerun()


def _bullet_list_path(section_key: str, item_idx: int) -> str:
    return f"{section_key}[{item_idx}].bullets"


def _apply_bullet_action(section_key: str, item_idx: int, bullet_idx: int, action: str) -> None:
    bullets = st.session_state.tailor_data.get(section_key, [])[item_idx].get("bullets", [])
    list_path = _bullet_list_path(section_key, item_idx)
    if action == "insert":
        insert_path = f"{list_path}[{bullet_idx + 1}]"
        if _apply_resume_patch([{"op": "add", "path": insert_path, "value": ""}], label="插入 bullet"):
            _open_inline_edit(insert_path, "")
            st.rerun()
    elif action == "delete":
        if len(bullets) <= 1:
            alert_warning("至少保留 1 条 bullet")
            return
        if _apply_resume_patch(
            [{"op": "remove", "path": f"{list_path}[{bullet_idx}]"}],
            label="删除 bullet",
        ):
            st.rerun()
    elif action in ("up", "down"):
        target_idx = bullet_idx - 1 if action == "up" else bullet_idx + 1
        if target_idx < 0 or target_idx >= len(bullets):
            return
        order = list(range(len(bullets)))
        order[bullet_idx], order[target_idx] = order[target_idx], order[bullet_idx]
        if _apply_resume_patch(
            [{"op": "reorder", "path": list_path, "value": order}],
            label="调整 bullet 顺序",
        ):
            st.rerun()


def _render_bullet(section_key: str, item_idx: int, bullet_idx: int, bullet: str) -> None:
    path = f"{_bullet_list_path(section_key, item_idx)}[{bullet_idx}]"
    if st.session_state.get(_path_state_key("edit_mode", path)):
        _render_edit_pane(path, bullet, height=86)
        return
    text_col, edit_col, add_col, up_col, down_col, del_col = st.columns([8, 0.55, 0.55, 0.55, 0.55, 0.55])
    with text_col:
        st.markdown(
            f'<div class="cv-bullet-read">• {resume_canvas.rich_text_html(bullet)}</div>',
            unsafe_allow_html=True,
        )
    with edit_col:
        if st.button("改", key=_path_state_key("edit", path), use_container_width=True, help="编辑这一条"):
            _open_inline_edit(path, bullet)
            st.rerun()
    with add_col:
        if st.button("加", key=_path_state_key("insert", path), use_container_width=True, help="在下方插入一条"):
            _apply_bullet_action(section_key, item_idx, bullet_idx, "insert")
    with up_col:
        if st.button("↑", key=_path_state_key("up", path), use_container_width=True, help="上移"):
            _apply_bullet_action(section_key, item_idx, bullet_idx, "up")
    with down_col:
        if st.button("↓", key=_path_state_key("down", path), use_container_width=True, help="下移"):
            _apply_bullet_action(section_key, item_idx, bullet_idx, "down")
    with del_col:
        if st.button("×", key=_path_state_key("delete", path), use_container_width=True, help="删除这一条"):
            _apply_bullet_action(section_key, item_idx, bullet_idx, "delete")


def _render_experience_canvas(section_key: str, title: str) -> None:
    data = st.session_state.tailor_data
    items = data.get(section_key) or []
    if not items:
        return
    resume_canvas.render_section(title)
    for item_idx, item in enumerate(items):
        st.markdown('<div class="cv-item">', unsafe_allow_html=True)
        role_path = f"{section_key}[{item_idx}].role"
        if st.session_state.get(_path_state_key("edit_mode", role_path)):
            resume_canvas.render_item_header(item.get("company", ""), "", item.get("date", ""))
            _render_edit_pane(role_path, item.get("role", ""), height=70)
        else:
            head_col, edit_col = st.columns([8, 1])
            with head_col:
                resume_canvas.render_item_header(
                    item.get("company", ""),
                    item.get("role", ""),
                    item.get("date", ""),
                )
            with edit_col:
                if st.button("编辑", key=_path_state_key("edit", role_path), use_container_width=True):
                    _open_inline_edit(role_path, item.get("role", ""))
                    st.rerun()
        for bullet_idx, bullet in enumerate(item.get("bullets") or []):
            _render_bullet(section_key, item_idx, bullet_idx, bullet)
        st.markdown("</div>", unsafe_allow_html=True)


def _render_skills_canvas() -> None:
    skills = st.session_state.tailor_data.get("skills") or []
    if not skills:
        return
    resume_canvas.render_section("技能证书")
    for skill_idx, skill in enumerate(skills):
        label_path = f"skills[{skill_idx}].label"
        text_path = f"skills[{skill_idx}].text"
        if (
            st.session_state.get(_path_state_key("edit_mode", label_path))
            or st.session_state.get(_path_state_key("edit_mode", text_path))
        ):
            label_key = _path_state_key("edit_draft", label_path)
            text_key = _path_state_key("edit_draft", text_path)
            st.session_state.setdefault(label_key, skill.get("label", ""))
            st.session_state.setdefault(text_key, skill.get("text", ""))
            st.markdown('<div class="cv-edit-pane">', unsafe_allow_html=True)
            c1, c2 = st.columns([1.2, 4])
            with c1:
                new_label = st.text_input("技能分类", key=label_key, label_visibility="collapsed")
            with c2:
                new_text = st.text_input("技能内容", key=text_key, label_visibility="collapsed")
            save_col, cancel_col, _ = st.columns([2, 2, 1])
            with save_col:
                if st.button("保存", key=f"save_skill_{skill_idx}", use_container_width=True):
                    patch = [
                        {"op": "replace", "path": label_path, "value": new_label},
                        {"op": "replace", "path": text_path, "value": new_text},
                    ]
                    if _apply_resume_patch(patch, label="手动编辑技能"):
                        _close_inline_edit(label_path)
                        _close_inline_edit(text_path)
                        st.rerun()
            with cancel_col:
                if st.button("取消", key=f"cancel_skill_{skill_idx}", use_container_width=True):
                    _close_inline_edit(label_path)
                    _close_inline_edit(text_path)
                    st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            row_col, edit_col = st.columns([8, 1])
            with row_col:
                st.markdown(
                    '<div class="cv-skill-row">'
                    f'<span class="cv-skill-label">{resume_canvas.esc(skill.get("label", ""))}</span>'
                    f'{resume_canvas.rich_text_html(skill.get("text", ""))}'
                    '</div>',
                    unsafe_allow_html=True,
                )
            with edit_col:
                if st.button("编辑", key=f"edit_skill_{skill_idx}", use_container_width=True):
                    _open_inline_edit(label_path, skill.get("label", ""))
                    _open_inline_edit(text_path, skill.get("text", ""))
                    st.rerun()


def _render_education_canvas() -> None:
    education = st.session_state.tailor_data.get("education") or []
    if not education:
        return
    resume_canvas.render_section("教育背景")
    for edu in education:
        school = resume_canvas.esc(edu.get("school", ""))
        major = resume_canvas.esc(edu.get("major", ""))
        date = resume_canvas.esc(edu.get("date", ""))
        st.markdown(
            f'<div class="cv-edu-row"><span><b>{school}</b> · {major}</span><span class="cv-item-meta">{date}</span></div>',
            unsafe_allow_html=True,
        )
        for bullet in edu.get("bullets") or []:
            st.markdown(
                f'<div class="cv-bullet-read">• {resume_canvas.rich_text_html(bullet)}</div>',
                unsafe_allow_html=True,
            )


def _render_resume_canvas_editor() -> None:
    """A4 document canvas: read mode first, edit mode only after an explicit click."""
    data = st.session_state.tailor_data
    meta = st.session_state.tailor_meta

    if meta:
        alert_info(
            f"**匹配度 {meta.get('match_score', '?')}** · "
            f"{meta.get('change_notes', '')}"
        )

    resume_canvas.render_canvas_css()

    with st.container(border=True):
        st.markdown('<span class="cos-canvas-anchor"></span>', unsafe_allow_html=True)
        _render_basics_editor(data)
        resume_canvas.render_basics_header(data.setdefault("basics", {}), data.get("education"))

        resume_canvas.render_section("个人总结")
        _render_editable_block("profile", data.get("profile", ""), block_class="cv-bullet-read", height=112)
        if st.button("仅重写个人总结", key="rw_profile"):
            if not st.session_state.tailor_jd:
                alert_warning("请先在左栏粘贴 JD")
            else:
                with st.spinner("重写中..."):
                    try:
                        intent = meta.get("jd_intent") or resume_tailor.extract_jd_intent(
                            st.session_state.tailor_jd
                        )
                        _push_undo_snapshot(label="重写个人总结")
                        data["profile"] = resume_tailor.rewrite_section(data["profile"], intent)
                        _clear_tailor_preview_cache()
                        st.rerun()
                    except AIError as e:
                        alert_danger(str(e))

        _render_experience_canvas("projects", "项目经历")
        _render_experience_canvas("internships", "实习经历")
        _render_skills_canvas()
        _render_education_canvas()


# ═══ 右栏：简历内容 + PDF 预览 + 深度评估 ═══
with col_right:
    tab_canvas, tab_pdf, tab_deep = st.tabs(["简历内容", "PDF 预览", "深度评估"])

with tab_deep:
    from services import resume_evaluator
    if "eval_data" not in st.session_state:
        st.session_state.eval_data = None

    c_btn, c_cache = st.columns([1, 1])
    with c_btn:
        run_eval = st.button("深度评估", type="primary", use_container_width=True,
                             disabled=not st.session_state.tailor_jd.strip())
    with c_cache:
        if st.button("读缓存", use_container_width=True,
                     disabled=not st.session_state.tailor_jd.strip()):
            cached = resume_evaluator.load_latest_for_jd(st.session_state.tailor_jd)
            if cached:
                st.session_state.eval_data = cached
                st.toast("已加载上次评估")
            else:
                alert_warning("此 JD 无历史评估")

    if run_eval:
        with st.spinner("跑深度评估中..."):
            try:
                ev = resume_evaluator.deep_evaluate(
                    st.session_state.tailor_jd,
                    st.session_state.tailor_data,
                    st.session_state.tailor_meta.get("jd_intent"),
                )
                st.session_state.eval_data = ev
                resume_evaluator.save_evaluation(
                    st.session_state.tailor_jd,
                    ev,
                    target_role=st.session_state.tailor_data["basics"].get("target_role", ""),
                )
                st.toast("评估完成 + 已落库")
            except Exception as e:
                alert_danger(f"评估失败：{e}")

    ev = st.session_state.eval_data
    if ev:
        score = ev.get("overall_score", 0)
        verdict = ev.get("one_line_verdict", "")
        score_hero_card(score, verdict=verdict)

        # A 匹配分析
        with st.expander("A · 匹配分析（逐条 JD 要求）", expanded=True):
            sa = ev.get("section_a_match", {})
            for req in sa.get("requirements", []):
                icon = {"match": "OK", "weak": "弱", "miss": "缺"}.get(req.get("status"), "?")
                st.markdown(f"{icon} **{req.get('requirement', '')}**")
                if req.get("evidence"):
                    st.caption(f"证据：「{req['evidence']}」 · {req.get('note', '')}")
                elif req.get("note"):
                    st.caption(req["note"])

        # B Gap
        with st.expander("B · Gap 分析"):
            sb = ev.get("section_b_gap", {})
            if sb.get("critical_gaps"):
                st.markdown("**硬缺口**")
                for g in sb["critical_gaps"]:
                    st.markdown(f"- {g}")
            if sb.get("soft_gaps"):
                st.markdown("**软缺口**")
                for g in sb["soft_gaps"]:
                    st.markdown(f"- {g}")
            if sb.get("bridgeable"):
                st.markdown("**可弥补**")
                for g in sb["bridgeable"]:
                    st.markdown(f"- {g}")

        # C 改写建议
        with st.expander("C · 简历改写建议"):
            sc = ev.get("section_c_resume_advice", {})
            if sc.get("strengthen"):
                st.markdown("**强化**")
                for s in sc["strengthen"]:
                    st.markdown(f"- {s}")
            if sc.get("downplay"):
                st.markdown("**弱化**")
                for s in sc["downplay"]:
                    st.markdown(f"- {s}")
            if sc.get("add_keywords"):
                st.markdown("**补关键词**: " + "、".join(sc["add_keywords"]))

        # D 预测题
        with st.expander("D · 预测面试题（5 道）"):
            sd = ev.get("section_d_interview_qa", {})
            for i, q in enumerate(sd.get("questions", []), 1):
                st.markdown(f"**Q{i}. {q.get('q','')}**")
                if q.get("why"):
                    st.caption(f"问法原因：{q['why']}")
                if q.get("answer_hint"):
                    st.markdown(f"> 回答要点：{q['answer_hint']}")
                st.markdown("")

        # E 策略
        with st.expander("E · 投递策略"):
            se = ev.get("section_e_strategy", {})
            rows = [
                ("推荐渠道", se.get("recommended_channel")),
                ("最佳时机", se.get("best_timing")),
                ("竞争激烈度", se.get("competition_level")),
                ("内推角度", se.get("referral_angle")),
                ("风险提示", se.get("risk_warning")),
            ]
            for k, v in rows:
                if v:
                    st.markdown(f"**{k}**：{v}")
    else:
        st.caption("先在左栏输入 JD 并生成定制版，再点「深度评估」")


with tab_canvas:
    alert_info("在线简历画布：右侧内容可直接编辑，PDF / DOCX 会按同一份内容生成。原简历 PDF 只作为参照，不作为编辑源。")
    _orig_pdf = st.session_state.get("uploaded_pdf_bytes")
    if _orig_pdf:
        import base64 as _b64
        with st.expander(f"你上传的原简历 PDF（{len(_orig_pdf):,} 字节）· 对照参考", expanded=False):
            _b64_s = _b64.b64encode(_orig_pdf).decode()
            st.markdown(
                f'<iframe src="data:application/pdf;base64,{_b64_s}" '
                f'width="100%" height="500px" style="border:1px solid rgba(29,29,31,0.08);border-radius:14px"></iframe>',
                unsafe_allow_html=True,
            )
    else:
        st.caption("（没上传过原 PDF · 去「主简历 → 上传文件」上传后，这里可以看到原件对照）")

    st.markdown("##### 在线简历画布")
    _render_resume_canvas_editor()

with tab_pdf:
    _render_pdf_preview_block(thumbnail=False)
